from typing import Dict, List, Optional
import os
from docx import Document
from pptx import Presentation
from ..services.openai_service import OpenAIService
from ..utils.file_handler import FileHandler
from ..utils.doc_converter import DocumentConverter

class DocumentGenerator:
    def __init__(self):
        self.openai_service = OpenAIService()
        self.file_handler = FileHandler()
        self.doc_converter = DocumentConverter()
        
    async def generate_training_doc(self, requirements: Dict, template_path: Optional[str] = None) -> str:
        """生成完整的培训文档"""
        try:
            # 生成大纲
            outline = await self.openai_service.generate_outline(requirements)
            
            # 解析大纲获取章节
            sections = self._parse_outline(outline)
            
            # 为每个章节生成内容
            content = {}
            for section in sections:
                section_content = await self.openai_service.generate_section_content(
                    section['title'],
                    requirements
                )
                content[section['title']] = section_content
            
            # 使用模板生成文档
            doc_path = self._create_document(content, template_path)
            
            return doc_path
            
        except Exception as e:
            raise Exception(f"Document generation failed: {str(e)}")
    
    def _parse_outline(self, outline: str) -> List[Dict]:
        """解析大纲文本，提取章节结构"""
        sections = []
        # 实现大纲解析逻辑
        return sections
    
    def _create_document(self, content: Dict, template_path: Optional[str]) -> str:
        """根据内容创建文档"""
        if template_path:
            doc = Document(template_path)
        else:
            doc = Document()
            
        # 添加内容到文档
        for section_title, section_content in content.items():
            doc.add_heading(section_title, level=1)
            doc.add_paragraph(section_content)
            
        # 保存文档
        output_path = os.path.join(Config.UPLOAD_FOLDER, 'generated_docs', 'training.docx')
        doc.save(output_path)
        
        return output_path
    
    def convert_document(self, doc_path: str, target_format: str) -> str:
        """转换文档格式"""
        return self.doc_converter.convert(doc_path, target_format) 